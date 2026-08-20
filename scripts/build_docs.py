from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(100, 106, 116)
LIGHT_FILL = "F2F4F7"
BLUE_FILL = "E8EEF5"


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent=120):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure_document(doc: Document, short_name: str):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(35, 38, 43)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run(f"Evidence-to-Alpha | {short_name}")
    set_run_font(run, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("v0.1.0 MVP | Internal working document")
    set_run_font(run, size=9, color=MUTED)


def add_title_block(doc, title, subtitle, status):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    set_run_font(run, size=23, color=RGBColor(0, 0, 0), bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(subtitle)
    set_run_font(run, size=13, color=RGBColor(70, 74, 80))
    for label, value in (("项目", "Evidence-to-Alpha Trading System"), ("版本", "v0.1.0 MVP"), ("日期", "2026-08-20"), ("状态", status)):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        label_run = p.add_run(f"{label}: ")
        set_run_font(label_run, size=10.5, color=RGBColor(35, 38, 43), bold=True)
        value_run = p.add_run(value)
        set_run_font(value_run, size=10.5, color=RGBColor(35, 38, 43))
    rule = doc.add_paragraph()
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D7DBE2")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    rule.paragraph_format.space_after = Pt(10)


def add_para(doc, text, bold_prefix=None, italic=False):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_run_font(first, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        set_run_font(rest, italic=italic)
    else:
        run = p.add_run(text)
        set_run_font(run, italic=italic)
    return p


def add_bullets(doc, items: Iterable[str], numbered=False):
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        p = doc.add_paragraph(style=style)
        run = p.add_run(item)
        set_run_font(run)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, BLUE_FILL)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header)
        set_run_font(run, size=10, color=DARK_BLUE, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            p = cells[index].paragraphs[0]
            run = p.add_run(str(value))
            set_run_font(run, size=9.5)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def build_selection():
    doc = Document()
    configure_document(doc, "开源选型与 MVP 建议")
    add_title_block(doc, "开源项目选型与 MVP 建议", "Evidence-to-Alpha Trading System", "选型结论已冻结；外部上线仍需授权")
    add_heading(doc, "执行摘要")
    add_para(doc, "结论：自己开发一个薄集成项目，不直接使用或 fork 完整交易引擎。核心差异是 point-in-time 证据链、事件版本不可覆盖、事件与因子信号叠加、逐笔 Paper OMS 追溯和研究门禁；这些约束无法从单一通用交易引擎直接获得。")
    add_para(doc, "建议复用成熟项目的局部能力或接口思想：Qlib 的研究记录、Zipline Reloaded 的交易日历和事件循环、NautilusTrader/LEAN 的订单模型设计。vectorbt 具备研究速度优势，但当前仓库含 Commons Clause，不作为基础依赖。")
    add_heading(doc, "候选项目对比")
    add_table(doc, ["项目", "维护/许可证", "部署复杂度", "可复用能力", "决策"], [
        ["Microsoft Qlib", "活跃；MIT；2026-07-23 推送；PyPI 0.9.7", "中高；依赖和数据准备较多", "数据集、工作流、记录器、组合分析", "参考，不 fork"],
        ["Zipline Reloaded", "活跃；Apache-2.0；2026-01-06 推送；PyPI 3.1.1", "中；pip/conda 可装，bundle 有门槛", "交易日历、事件循环、绩效", "后续可选适配器"],
        ["NautilusTrader", "高活跃；LGPL-3.0；2026-08-20 推送；PyPI 1.231.0", "高；Rust/Python 3.12+", "确定性事件驱动、订单/成交/持仓", "实盘阶段再评估"],
        ["QuantConnect LEAN", "高活跃；Apache-2.0；2026-08-19 推送；CLI 1.0.228", "高；Docker/.NET/数据约定", "多资产回测和 Paper/Live", "功能过重"],
        ["vectorbt", "活跃；2026-08-02 推送；Commons Clause", "低到中；pip/Docker", "向量化扫描、事件消融", "许可证风险，排除"],
    ], [1450, 2200, 1650, 2600, 1460])
    add_heading(doc, "最简单 MVP")
    add_bullets(doc, [
        "输入五类文件：事件版本 JSON、证据 JSON、实体映射 CSV、日线行情 CSV、因子基线权重 CSV。",
        "处理六件事：as-of 版本选择、实体映射、事件信号、1/3/5/20 日事件研究、有限幅度 overlay、T+1 Paper OMS。",
        "输出报告、信号、订单、成交、事件研究 CSV、SQLite 账本、自动验证结果和只读 HTTP API。",
        "不做新闻前端、不接券商、不读取真实凭据、不训练复杂 ML、不把样例数据包装成实盘结论。",
    ])
    add_heading(doc, "选型决策依据")
    add_table(doc, ["选项", "优点", "主要风险", "最终判断"], [
        ["直接使用", "上线快、通用能力完整", "无法满足证据/版本/归因契约；引入过多边界", "不选"],
        ["基于现有项目改", "可复用既有因子平台", "当前工作区为空；跨项目耦合和发布边界不清", "不作为本仓库方案"],
        ["自己开发薄集成", "边界清晰、可审计、部署简单", "需自行维护契约和 OMS", "推荐"],
    ], [1900, 2450, 3000, 2010])
    add_heading(doc, "证据与限制")
    add_para(doc, "调研结论基于 2026-08-20 读取的 GitHub 公共仓库元数据、项目 README/LICENSE 和 PyPI 元数据。维护活跃不等于适合本项目；许可证、数据源、API 稳定性和商业化限制必须在正式发布前再次复核。")
    add_bullets(doc, [
        "https://github.com/microsoft/qlib",
        "https://github.com/stefan-jansen/zipline-reloaded",
        "https://github.com/nautechsystems/nautilus_trader",
        "https://github.com/QuantConnect/Lean",
        "https://github.com/polakowo/vectorbt",
    ])
    return doc


def build_prd():
    doc = Document()
    configure_document(doc, "产品需求文档")
    add_title_block(doc, "产品需求文档", "证据驱动的事件增强交易系统", "需求已冻结；MVP 目标为研究与纸面交易")
    add_heading(doc, "1. 产品目标")
    add_para(doc, "在严格使用当时可得新闻证据的条件下，验证事件信号能否为既有多因子策略带来稳定、可解释、扣除成本后的增量信息。系统允许 PROMOTE、REJECT、INCONCLUSIVE 三种结论，后两种结论同样必须保留完整证据链。")
    add_heading(doc, "2. 用户与场景")
    add_bullets(doc, [
        "量化研究员导入新闻系统只读快照，生成事件研究和信号。",
        "组合研究员将事件信号作为有限幅度 overlay 叠加到因子基线。",
        "验证人员检查泄漏、成本、安慰剂、延迟和样本外证据。",
        "研究负责人将任一订单追溯到事件版本、原始证据、映射和因子版本。",
    ])
    add_heading(doc, "3. 范围与非范围")
    add_table(doc, ["范围", "MVP 内容"], [
        ["包含", "事件契约、版本选择、实体映射、信号、事件研究、组合叠加、Paper OMS、验证、审计、CLI、只读 API"],
        ["不包含", "新闻采集/前端、券商连接、真实资金、生产凭据、复杂 ML、实时撮合、高可用集群"],
    ], [2100, 7260])
    add_heading(doc, "4. 功能需求")
    requirements = [
        ("FR-01 事件数据适配", "读取 event_id、event_version、published_at、observed_at、asof、影响主体和证据 ID；同一事件版本不可被不同内容覆盖。"),
        ("FR-02 Point-in-time 选择", "运行只选择 asof 不晚于截止时间的最新版本；未来版本不得影响历史运行。"),
        ("FR-03 实体映射", "映射公司或行业实体到股票代码、行业和影响倍数；失败映射必须披露。"),
        ("FR-04 信号生成", "方向、可信度、新颖性、冲突和时间衰减共同决定信号；信号携带配置哈希、事件版本和证据 ID。"),
        ("FR-05 事件研究", "计算 1、3、5、20 个交易日绝对和相对基准收益；数据不足必须标记。"),
        ("FR-06 信号融合", "提供因子基线、事件基线和 overlay，并受单票调整、换手和单票权重限制。"),
        ("FR-07 Paper OMS", "下一可交易日开盘按滑点和手续费成交，维护现金、持仓和 PnL，完成账本闭合。"),
        ("FR-08 研究验证", "执行时间泄漏、证据完整性、版本不可变、T+1、成本翻倍、延迟一日、安慰剂、基线和样本量检查。"),
        ("FR-09 追溯输出", "输出研究结论、门禁、事件研究、信号、订单、成交和 SQLite 账本。"),
        ("FR-10 只读服务", "提供健康检查、最新运行、信号、订单、成交和事件研究 JSON 端点。"),
    ]
    add_table(doc, ["编号", "需求"], [[name, detail] for name, detail in requirements], [2200, 7160])
    add_heading(doc, "5. 非功能需求")
    add_bullets(doc, [
        "确定性：相同输入和配置产生相同配置哈希、信号和订单。",
        "安全：不读取券商凭据，不提供写交易端点。",
        "可复现：运行报告记录输入摘要、配置和版本。",
        "可部署：Python 3.11+ 直接运行，提供 Docker 配置。",
        "可测试：核心领域逻辑由标准库单元测试覆盖。",
        "可审计：重要结论区分已验证事实、推断、未知和决策。",
    ])
    add_heading(doc, "6. 验收标准")
    add_table(doc, ["门禁", "通过标准"], [
        ["时间与版本", "未来版本不能进入历史运行；observed_at 不早于 published_at。"],
        ["证据链", "所有信号引用的 evidence_id 存在；事件版本和配置哈希保留。"],
        ["交易时序", "每笔成交日期严格晚于信号截止日期。"],
        ["组合限制", "权重和为 1，overlay、换手和单票权重不超限。"],
        ["账本闭合", "现金、持仓、市值、费用和总权益闭合到 0.01。"],
        ["研究结论", "基线和增强独立报告；样本不足返回 INCONCLUSIVE。"],
        ["交付质量", "测试、验证证据和文档渲染均通过；生产发布仍需人工授权。"],
    ], [2100, 7260])
    add_heading(doc, "7. 研究结论边界")
    add_para(doc, "本系统是研究和纸面交易系统。没有真实成交数据、稳定样本外表现和独立风险验证之前，不得称为有效实盘交易系统。")
    return doc


def build_dev():
    doc = Document()
    configure_document(doc, "开发与部署文档")
    add_title_block(doc, "开发与部署文档", "Evidence-to-Alpha Trading System v0.1.0", "实现候选已完成；本地发布候选待冻结")
    add_heading(doc, "1. 架构决策")
    add_para(doc, "采用 Python 标准库实现薄集成服务，不 fork 完整交易引擎。核心价值是证据和版本契约，而不是重新实现通用量化平台。Zipline Reloaded 可作为后续回测适配器；NautilusTrader 或 LEAN 只在真实执行评估阶段再考虑。")
    add_heading(doc, "2. 组件与责任")
    add_table(doc, ["组件", "责任", "关键证据"], [
        ["contracts", "解析并验证事件、证据、映射、行情和权重", "时间字段、版本、输入摘要"],
        ["store", "SQLite 事件版本和运行账本", "不可变键、run_id、artifact"],
        ["signals", "门控、衰减、配置哈希、逐条 lineage", "signal_id、event_ref、evidence_id"],
        ["study", "1/3/5/20 日异常收益", "窗口状态、基准收益"],
        ["portfolio", "基线、事件、overlay 和约束", "权重和、单票、换手"],
        ["oms", "T+1 订单、滑点、手续费、持仓与闭合", "order、fill、reconciliation"],
        ["validation", "泄漏、证据、压力和结论门禁", "audit.json"],
        ["api/cli", "只读服务和可复现命令入口", "health、report、demo"],
    ], [1500, 4750, 3110])
    add_heading(doc, "3. 数据与版本契约")
    add_bullets(doc, [
        "事件唯一键为 event_id + event_version；published_at、observed_at、asof 为带时区 ISO-8601。",
        "published_at <= observed_at <= asof；运行截止时间决定可见版本。",
        "订单和成交记录 run_id、ticker、side、quantity、factor_version、signal_ids、event_refs、evidence_ids。",
        "没有事件调整的因子基线订单允许事件引用为空，但 factor_version 不可为空。",
    ])
    add_heading(doc, "4. 核心流程")
    add_bullets(doc, [
        "读取并验证输入，登记所有事件版本到 SQLite。",
        "按截止时间选择每个事件的最新可见版本。",
        "映射实体并生成带配置哈希的事件信号。",
        "计算事件研究和异常收益，生成三条组合路径。",
        "下一交易日生成纸面成交，执行成本和账本闭合。",
        "运行自动验证，输出 PROMOTE / REJECT / INCONCLUSIVE。",
        "写出 JSON、CSV、SQLite 和审计证据。",
    ], numbered=True)
    add_heading(doc, "5. 运行命令")
    add_table(doc, ["操作", "命令"], [
        ["安装", "python -m pip install -e ."],
        ["演示", "python -m evidence_alpha demo --output-dir artifacts/demo"],
        ["测试", "python -m unittest discover -s tests -v"],
        ["验证", "python -m evidence_alpha verify --artifact-dir artifacts/demo"],
        ["服务", "python -m evidence_alpha serve --artifact-dir artifacts/demo --port 8080"],
        ["Docker", "docker compose up --build"],
    ], [1700, 7660])
    add_heading(doc, "6. API")
    add_table(doc, ["端点", "用途"], [
        ["GET /health", "服务状态和报告是否就绪"],
        ["GET /api/v1/runs/latest", "最新运行报告、结论和门禁"],
        ["GET /api/v1/runs/latest/signals", "事件信号及配置/证据追溯"],
        ["GET /api/v1/runs/latest/orders", "目标权重到订单追溯"],
        ["GET /api/v1/runs/latest/fills", "订单到成交追溯"],
        ["GET /api/v1/runs/latest/event-study", "事件研究 CSV 的 JSON 表示"],
    ], [3300, 6060])
    add_heading(doc, "7. 测试与发布门禁")
    add_para(doc, "当前测试覆盖时间契约、版本选择、不可变输入、信号追溯、组合限制、T+1、账本闭合和端到端 demo。发布前还要运行源码编译、自动 verifier、API 冒烟和文档渲染检查。")
    add_bullets(doc, [
        "自动化测试是近端证据，不替代独立 Verifier。",
        "最后一次文件编辑后只执行 VERIFY -> SEAL -> CLEANUP -> FINAL。",
        "Git 合并、远程仓库、云主机、TLS、凭据、真实数据库、生产发布和回滚仍是单独授权门禁。",
    ])
    add_heading(doc, "8. 已知限制与上线前置条件")
    add_para(doc, "本地发布候选可运行，但工作区没有远程 Git 地址、云账号、域名、证书或生产数据授权。因此本轮完成本地 API/Docker 部署与证据封存；外部互联网生产上线必须由发布负责人补齐目标环境和授权后执行。")
    return doc


def save(doc: Document, name: str):
    OUT.mkdir(parents=True, exist_ok=True)
    path = OUT / name
    doc.save(path)
    return path


if __name__ == "__main__":
    paths = [
        save(build_selection(), "01_开源选型与MVP建议.docx"),
        save(build_prd(), "02_产品需求文档_PRD.docx"),
        save(build_dev(), "03_开发与部署文档.docx"),
    ]
    for path in paths:
        print(path)

