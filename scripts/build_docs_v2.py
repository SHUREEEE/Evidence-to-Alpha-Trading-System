from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

import build_docs as base


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables"
INK = RGBColor(35, 38, 43)
MUTED = RGBColor(100, 106, 116)
RULE = "D7DBE2"


def configure(doc: Document, short_name: str) -> None:
    base.BLUE_FILL = "F2F4F7"
    base.configure_document(doc, short_name)
    footer = doc.sections[0].footer.paragraphs[0]
    footer.clear()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("v0.5.0 | Internal research document | Page ")
    base.set_run_font(run, size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run_element = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "646A74")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    props.extend([color, size])
    text = OxmlElement("w:t")
    text.text = "1"
    run_element.extend([props, text])
    field.append(run_element)
    footer._p.append(field)


def title_block(doc: Document, title: str, subtitle: str, status: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    base.set_run_font(p.add_run(title), size=23, color=RGBColor(0, 0, 0), bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    base.set_run_font(p.add_run(subtitle), size=13, color=RGBColor(70, 74, 80))
    rows = (
        ("项目", "Evidence-to-Alpha Trading System"),
        ("版本", "v0.5.0 Real-Data Readiness Preflight"),
        ("日期", "2026-08-21"),
        ("状态", status),
    )
    for label, value in rows:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        base.set_run_font(p.add_run(f"{label}: "), size=10.5, color=INK, bold=True)
        base.set_run_font(p.add_run(value), size=10.5, color=INK)
    rule = doc.add_paragraph()
    p_pr = rule._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), RULE)
    borders.append(bottom)
    p_pr.append(borders)
    rule.paragraph_format.space_after = Pt(10)


def para(doc: Document, text: str, lead: str | None = None) -> None:
    base.add_para(doc, text, bold_prefix=lead)


def bullets(doc: Document, items: list[str], numbered: bool = False) -> None:
    if not numbered:
        base.add_bullets(doc, items)
        return

    numbering = doc.part.numbering_part.element
    style_num_id = doc.styles["List Number"]._element.pPr.numPr.numId.val
    style_num = numbering.xpath(f'./w:num[@w:numId="{style_num_id}"]')[0]
    abstract_num_id = style_num.xpath("./w:abstractNumId")[0].get(qn("w:val"))
    next_num_id = max(int(item.get(qn("w:numId"))) for item in numbering.findall(qn("w:num"))) + 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(next_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_num_id)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.extend([abstract, override])
    numbering.append(num)

    for item in items:
        paragraph = doc.add_paragraph(style="List Number")
        base.set_run_font(paragraph.add_run(item))
        num_pr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
        num_pr.get_or_add_ilvl().val = 0
        num_pr.get_or_add_numId().val = next_num_id


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    value = base.add_table(doc, headers, rows, widths)
    for index, row in enumerate(value.rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        if index == 0:
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)


def build_selection() -> Document:
    doc = Document()
    configure(doc, "开源选型与 MVP 建议")
    title_block(doc, "开源项目选型与 MVP 建议", "三系统交易研究闭环", "选型保持不变；集成独立验证已实现；Live 阻断")

    base.add_heading(doc, "执行结论")
    para(doc, "建议：基于现有两个项目增加薄集成层。不要直接引入第三套完整交易引擎，也不要重写多因子平台。", "建议：")
    para(doc, "新项目把新闻证据转换为 point-in-time 事件 Alpha，并在 multi-factor-alpha-platform 的 V4 组合控制之前注入有限幅度权重调整。因子平台继续拥有因子、V4 风控、T+1 回测与归因；新闻系统继续拥有采集、事件、证据和影响分析。")
    para(doc, "闭环：新闻证据 -> 事件版本 -> 事件 Alpha -> 因子 + 事件预 V4 权重 -> V4 风控 -> T+1 回测/Paper OMS -> 归因反馈。")

    base.add_heading(doc, "三项目职责")
    table(doc, ["项目", "保留职责", "集成边界"], [
        ["News Claws / 127.0.0.1:8765", "新闻采集、事件聚类、来源证据、可信度、公司/行业影响、时间线", "只读 GET；逐版本导出，不读内部 SQLite"],
        ["Evidence-to-Alpha", "时间契约、事件信号、三路融合、追溯、Paper OMS、研究门禁", "编排层；不复制新闻前端和因子库"],
        ["multi-factor-alpha-platform", "V3 权重、V4 风控、T+1 回测、风险归因、发布门禁", "CSV/Parquet 文件契约；可选子进程验证"],
    ], [2450, 3800, 3110])

    base.add_heading(doc, "现有多因子项目评估")
    table(doc, ["维度", "直接验证结果"], [
        ["维护", "提交 9792ed27059b1179b39cca8fca2982fe22baf86e；2026-06-21 23:50:03 +08:00"],
        ["许可证", "MIT"],
        ["可复用", "run_backtest.py、run_v4_pipeline.py、V3 权重契约、V4 风控、归因和发布门禁"],
        ["部署", "中等：Python、pandas、pyarrow、cvxpy；大型权重和价格文件需在本地恢复"],
        ["公共仓库缺口", "v3_weights.parquet 与 data/processed/prices.parquet 未提交；sector map 可见"],
        ["Live 状态", "BLOCKED：真实 PB borrow feed 和 dry-run manifest 未满足"],
    ], [2100, 7260])
    para(doc, "本轮生成的预 V4 cache 被该仓库 production loader 接受，清单为 input_mode=prod、validation_state=PASS、cvxpy 路径通过；同时 borrow_feed_present=false。")

    base.add_heading(doc, "新闻系统评估")
    bullets(doc, [
        "只读端点已跑通：事件列表、逐版本详情、时间线、统计和来源列表。",
        "逐版本详情包含 claim/evidence、source URL、NVDA/TSM ticker、行业影响、可靠性、新颖性和不确定性。",
        "当前数据为 1 个 synthetic event、2 个版本；默认禁止进入研究，只有显式 opt-in 才能运行。",
        "discovered_at/created_at 早于 published_at 时采用最大时间作为 observed_at，宁可延后，不提前可见性。",
    ])

    base.add_heading(doc, "开源候选对比")
    table(doc, ["项目", "维护/许可", "部署", "可复用", "决策"], [
        ["Microsoft Qlib", "活跃；MIT", "中高", "工作流、记录器、组合分析", "参考，不做内核"],
        ["Zipline Reloaded", "活跃；Apache-2.0", "中", "日历、事件循环、绩效", "后续适配器"],
        ["NautilusTrader", "高活跃；LGPL-3.0", "高", "订单、成交、持仓", "实盘阶段再评估"],
        ["QuantConnect LEAN", "高活跃；Apache-2.0", "高", "多资产、Paper/Live、券商", "过重，不引入"],
        ["vectorbt", "活跃；Commons Clause", "低中", "参数扫描、事件消融", "许可证风险，排除"],
    ], [1760, 1750, 900, 3070, 1880])
    para(doc, "候选信息基于 2026-08-20 已记录的 GitHub、README、LICENSE 和 PyPI 调研。本轮 GitHub REST 公共 API 被共享出口限流，未把失败请求当作新证据。")

    base.add_heading(doc, "最简单 MVP")
    bullets(doc, [
        "从新闻 API 逐版本导出事件、证据、URL 和 ticker。",
        "从 V3 权重、行业映射和 adj_close 生成 as-of 因子基线。",
        "输出因子基线、事件基线、因子 + 事件三路权重。",
        "预 V4 delta 保持净敞口不变，并限制单票和总换手。",
        "生成 V4 cache，运行 V4 production loader 和三路 T+1 回测。",
        "生成 T+1 Paper 订单，逐笔携带 signal/event/evidence lineage。",
        "提供只读结果 API；没有任何券商或真实订单写入口。",
    ], numbered=True)

    doc.add_page_break()
    base.add_heading(doc, "验证状态")
    table(doc, ["分类", "结论", "证据/限制"], [
        ["已验证事实", "集成和验证接口可运行", "22/22 测试；integrate 已接入事件研究、独立 IS/OOS、滚动、placebo、延迟和成本翻倍门禁"],
        ["推断", "同契约真实数据可接入", "真实 V3 大文件尚未发现，仍需用户本地副本"],
        ["未知", "经济增量和实盘质量", "当前仅 2 个 synthetic 事件；真实 OOS、容量、borrow、成交尚未验证"],
        ["决策", "研究/Paper 可用", "当前 INCONCLUSIVE；Live BLOCKED"],
    ], [1600, 2150, 5610])

    base.add_heading(doc, "v0.4.0 集成验证更新")
    bullets(doc, [
        "独立验证器已从核心 run 路径接入真实三系统 integrate 路径。",
        "集成输出新增 event_study.csv、五个数值稳健性场景、independent_validation.json 和标准 audit.json。",
        "Synthetic manifest 会覆盖任何 real 声明并强制分类为 synthetic，不能 PROMOTE。",
        "验证运行 INT-3294801BE2C27699：13 个集成硬门禁全部通过，V4 loader 与三路回测均 PASS，标准 verify 通过。",
        "当前仅 1 个 synthetic 可见事件，5 日主窗口样本不足，决策仍为 INCONCLUSIVE；Live 继续 BLOCKED。",
    ])

    base.add_heading(doc, "最终选择")
    para(doc, "选择“基于现有项目改造并新增薄集成层”。这比直接使用通用引擎部署更简单，也比完全自研更尊重现有资产和责任边界。")
    return doc


def build_prd() -> Document:
    doc = Document()
    configure(doc, "产品需求文档")
    title_block(doc, "产品需求文档", "证据驱动的事件增强交易系统", "需求已冻结；v0.4.0 Integrated Validation")

    base.add_heading(doc, "1. 产品目标")
    para(doc, "在严格使用当时可得新闻证据的条件下，把新闻事件转换为可回测事件 Alpha，并与现有多因子组合融合，形成研究、V4 风控、T+1 回测、Paper OMS 和归因反馈闭环。")
    para(doc, "成功不是单次更高 Sharpe，而是回答：事件信号扣除成本并通过泄漏、样本外和独立验证后，是否给因子基线带来稳定增量。允许 PROMOTE、REJECT、INCONCLUSIVE。")

    base.add_heading(doc, "2. 用户与核心场景")
    bullets(doc, [
        "量化研究员：选择 as-of，运行三路比较和事件研究。",
        "组合研究员：检查预 V4 事件调整、换手和股票池重叠。",
        "风险/验证人员：核对 point-in-time、证据、V4、T+1 和账本门禁。",
        "研究负责人：从 Paper 订单追溯到 signal、event version、evidence 和 source URL。",
        "运维人员：启动只读结果服务，检查健康和最新报告。",
    ])

    base.add_heading(doc, "3. 系统边界")
    table(doc, ["系统", "所有权", "本项目交互"], [
        ["新闻系统", "采集、事件、证据、可靠性、影响", "只读 HTTP GET，逐版本导出"],
        ["Evidence-to-Alpha", "契约、信号、融合、追溯、Paper OMS、门禁", "核心编排"],
        ["多因子平台", "因子、V3 权重、V4、T+1、归因", "文件契约和可选外部验证"],
        ["券商/PB", "borrow、真实订单和成交", "MVP 不连接，保持阻断"],
    ], [2000, 3730, 3630])

    base.add_heading(doc, "4. 功能需求")
    table(doc, ["编号", "需求", "验收要点"], [
        ["FR-001", "新闻只读接入", "只调用 GET；不修改事件、来源、审核和数据库"],
        ["FR-002", "逐版本导出", "timeline 每个 version 通过 ?version=N 单独获取"],
        ["FR-003", "时间契约", "带时区；published <= observed <= asof；异常时间保守延后"],
        ["FR-004", "Synthetic 门禁", "默认拒绝；显式 opt-in 后仍不得 PROMOTE"],
        ["FR-005", "证据追溯", "signal 包含 event/version/evidence/config；URL 可查"],
        ["FR-006", "多因子输入", "长/宽 CSV/Parquet 权重；adj_close；sector map"],
        ["FR-007", "As-of/股票池", "权重不晚于 cutoff；重叠率低于门槛失败"],
        ["FR-008", "事件信号", "方向、可靠性、新颖性、冲突、衰减；失败映射披露"],
        ["FR-009", "三路组合", "因子、事件、融合独立输出，不只展示最好结果"],
        ["FR-010", "预 V4 融合", "delta 净和 0；单票和换手受限；在 V4 前执行"],
        ["FR-011", "V4 交接", "生成 v3_weights.parquet 与 v3_sector_map.csv"],
        ["FR-012", "外部回测", "三路由 run_backtest.py 接受；记录 return code/metrics"],
        ["FR-013", "Paper OMS", "T+1、side/qty/price/fee/lineage、会计闭合"],
        ["FR-014", "集成独立验证", "run 与 integrate 均按 observed_at 划分 IS/OOS；滚动折叠；泄漏/交叉分区硬拒绝"],
        ["FR-015", "稳健性与决策", "基线/placebo/延迟/成本翻倍；不足 INCONCLUSIVE；失败 REJECT"],
        ["FR-016", "只读服务", "health/report/signals/orders/event-study/independent-validation；POST 405"],
        ["FR-017", "可复现产物", "run ID、输入、配置、event_study.csv、audit.json、门禁和局限"],
        ["FR-018", "实盘阻断", "PB/真实样本/OOS/独立验证/授权缺一则 BLOCKED"],
    ], [1200, 3150, 5010])

    base.add_heading(doc, "5. 非功能需求")
    table(doc, ["属性", "要求"], [
        ["安全", "无券商凭据、无真实订单写路径、新闻侧只读"],
        ["确定性", "相同输入和配置产生相同 signal/order/run ID"],
        ["兼容性", "核心标准库；Parquet/V4 为 pandas/pyarrow 可选依赖"],
        ["可测试", "覆盖失败关闭、时间异常、版本、融合和外部契约"],
        ["可审计", "事实、推断、未知、决策分离；synthetic 指标带限制"],
        ["可运维", "健康检查、Docker、可配置 artifact directory"],
    ], [1800, 7560])

    base.add_heading(doc, "6. 研究运行")
    bullets(doc, [
        "检查新闻服务与因子输入。",
        "导出所有事件版本并选择 as-of 可见版本。",
        "生成信号、三路权重和预 V4 cache。",
        "运行 V4 production loader、三路回测和 Paper OMS。",
        "生成报告、审计产物和只读 API。",
    ], numbered=True)

    base.add_heading(doc, "7. 失败关闭")
    bullets(doc, [
        "Synthetic 未授权、证据缺失、版本冲突或股票池无重叠：立即失败。",
        "Parquet 引擎缺失：CSV 仍可写；请求 V4 时失败并披露原因。",
        "PB feed 缺失：研究继续，但 Live 永久保持 BLOCKED。",
        "样本不足：输出 INCONCLUSIVE，不继续调参或包装收益。",
    ])

    base.add_heading(doc, "8. 验收标准")
    table(doc, ["门禁", "通过标准", "当前状态"], [
        ["新闻联调", "真实 GET，所有版本导出", "PASS：1 event / 2 versions"],
        ["时间/版本", "未来版本不可见，异常时间不提前", "PASS"],
        ["Synthetic", "默认拒绝，显式 opt-in 仍不 PROMOTE", "PASS"],
        ["证据链", "事件订单可追溯到 evidence 和 URL", "PASS"],
        ["预 V4", "net delta=0；turnover<=0.08", "PASS"],
        ["外部接口", "V4 prod loader + 三路回测 return code 0", "PASS"],
        ["Paper OMS", "T+1 且闭合到 0.01", "PASS"],
        ["独立验证器", "IS/OOS、滚动、稳健性、失败关闭", "PASS：机制完成；样本不足"],
        ["集成验证产物", "integrate 写事件研究、五场景、独立验证和标准审计", "PASS：8 行 event study；5/5 场景数值化"],
        ["标准审计", "evidence-alpha verify 可验证 integrated artifact", "PASS：0 个硬失败"],
        ["自动测试", "全部通过且无 ResourceWarning", "PASS：22/22"],
        ["经济价值", "真实样本 + OOS + 独立验证", "INCONCLUSIVE"],
        ["Live", "PB borrow 等 P0 全部 READY", "BLOCKED"],
    ], [2100, 4300, 2960])

    base.add_heading(doc, "9. 后续路线")
    bullets(doc, [
        "接入真实 V3 权重、价格和至少 100 个真实事件版本。",
        "用真实样本重跑已实现的滚动 OOS、placebo、延迟和成本翻倍门禁。",
        "将 V4 后权重和归因回写为只读研究结果。",
        "完成 PB borrow feed、独立风险验证和连续 Paper 运行期。",
        "券商连接和实盘发布走单独授权。",
    ], numbered=True)
    return doc


def build_development() -> Document:
    doc = Document()
    configure(doc, "开发与部署文档")
    title_block(doc, "开发与部署文档", "Evidence-to-Alpha v0.5.0", "真实数据预检机制已验证；真实输入和 Live 仍阻断")

    base.add_heading(doc, "1. 架构决策")
    para(doc, "采用薄集成层：News Claws API -> NewsAdapter -> immutable EventSnapshot -> EventSignal -> pre-V4 Overlay -> multi-factor V4 -> T+1 backtest/Paper OMS -> IndependentValidation -> read-only API。")
    para(doc, "事件 Alpha 必须在 V4 之前进入。换手惩罚、no-trade band、行业净敞口和优化器仍由多因子平台负责，避免双重风控或约束顺序错误。")

    base.add_heading(doc, "2. 模块")
    table(doc, ["模块", "路径", "责任"], [
        ["NewsAdapter", "news_adapter.py", "GET、逐版本转换、时间/synthetic 门禁、证据导出"],
        ["MultiFactorAdapter", "multifactor_adapter.py", "CSV/Parquet、宽/长表、sector、V4 cache"],
        ["Integration", "integration.py", "重叠、三路组合、预 V4、事件研究、五场景、独立验证、标准审计、Paper OMS"],
        ["Contracts/Models", "contracts.py / models.py", "事件、证据、价格、权重类型与校验"],
        ["Signals", "signals.py", "门控、衰减和 lineage"],
        ["Independent Validation", "independent_validation.py", "按时间划分 IS/OOS、滚动折叠、稳健性门禁和三态决策"],
        ["API/CLI", "api.py / cli.py", "只读 artifact 服务、inspect-panel、readiness 和可复现命令"],
    ], [1800, 2750, 4810])

    base.add_heading(doc, "3. 新闻 API 契约")
    table(doc, ["端点", "用途"], [
        ["GET /api/v1/events", "分页事件列表"],
        ["GET /api/v1/events/{id}/timeline", "全部 version 编号与版本元数据"],
        ["GET /api/v1/events/{id}?version=N", "版本对应 claims、evidence、impacts 和 articles"],
    ], [4000, 5360])
    bullets(doc, [
        "只接受 data_version v1；缺少 data envelope 失败。",
        "published_at 为相关证据文章最早发布时间。",
        "observed_at 为相关 published/discovered/version.created 的最大值。",
        "Synthetic metadata 或 license 默认触发失败关闭。",
    ])

    base.add_heading(doc, "4. 多因子契约")
    table(doc, ["输入/输出", "支持格式", "关键字段"], [
        ["权重", "CSV/Parquet；长表或宽表", "date,ticker,weight 或 date + ticker columns"],
        ["行情", "CSV/Parquet；长表或宽表", "date,ticker,adj_close/total_return_index 或宽表"],
        ["行业", "CSV", "symbol/ticker,sector"],
        ["V4 cache", "Parquet + CSV", "v3_weights.parquet、v3_sector_map.csv"],
    ], [1900, 2950, 4510])

    base.add_heading(doc, "5. 融合算法")
    bullets(doc, [
        "按 ticker 聚合 as-of 可见事件信号。",
        "raw_delta = event_score * overlay_scale。",
        "横截面去均值，使 delta 净和为 0。",
        "应用单票绝对上限并重新分配残差。",
        "gross delta 超过 turnover cap 时整体缩放。",
        "fused_pre_v4 = factor_baseline + delta，然后交给 V4。",
    ], numbered=True)
    para(doc, "默认参数：overlay_scale=0.02，max_overlay_per_name=0.01，overlay_turnover_cap=0.08，cost_bps=5。")

    base.add_heading(doc, "6. 安装与运行")
    table(doc, ["操作", "命令"], [
        ["核心安装", "python -m pip install -e ."],
        ["Parquet/V4", "python -m pip install -e .[integrations]"],
        ["测试", "python -m unittest discover -s tests -v"],
        ["独立验证演示", "python -m evidence_alpha demo --output-dir artifacts/demo-v0.4"],
        ["新闻导出", "python -m evidence_alpha news-export --news-base-url http://127.0.0.1:8765 --output-dir artifacts/news"],
        ["市场输入检查", "python -m evidence_alpha inspect-panel --input <file> --kind <factor_weights|adjusted_prices>"],
        ["发布就绪检查", "python -m evidence_alpha readiness --artifact-dir <integrated> --pb-ingestion-manifest <manifest> ..."],
        ["结果服务", "python -m evidence_alpha serve --artifact-dir artifacts/integrated --host 127.0.0.1 --port 8080"],
        ["Docker", "docker compose up --build"],
    ], [1900, 7460])
    para(doc, "完整集成命令：")
    para(doc, "python -m evidence_alpha integrate --news-base-url http://127.0.0.1:8765 --factor-root <multi-factor-root> --asof <timezone-aware-timestamp> --benchmark SPY --data-classification unknown --minimum-event-count 30 --minimum-oos-events 10 --rolling-folds 3 --primary-window-days 5 --output-dir artifacts/integrated --run-factor-v4 --run-factor-backtests")
    para(doc, "公共仓库缺少大型二进制文件时，用 --weights、--sectors、--prices 指向本地真实产物；demo 新闻需要显式 --allow-synthetic-news。")

    base.add_heading(doc, "7. 产物")
    table(doc, ["路径", "内容"], [
        ["integration_report.json", "状态、三路比较、门禁、外部验证、live block"],
        ["integration_audit.json", "机器可读集成门禁"],
        ["audit.json", "标准门禁列表、三态决策、事实/推断/未知；供 verify 使用"],
        ["event_study.csv", "逐事件、ticker、1/3/5/20 日窗口及异常收益状态"],
        ["independent_validation.json", "IS/OOS 分区、滚动折叠、场景收益、门禁和三态结论"],
        ["signals.json / orders.json / fills.json", "信号、T+1 订单与成交证据链；兼容保留 paper_orders.json"],
        ["*_weights.csv", "因子、事件、融合三路回测输入"],
        ["v4_input_cache", "V4 production loader 输入"],
        ["factor_backtests / v4_output", "外部回测与 V4 清单"],
    ], [3300, 6060])

    base.add_heading(doc, "8. 测试和本轮证据")
    bullets(doc, [
        "63/63 unittest 和 8 个 subtest 通过；所有 warning 按错误处理后仍为零。",
        "独立验证覆盖正向 PROMOTE、负向 OOS REJECT、小样本 INCONCLUSIVE、泄漏、未知引用、坏行和非有限场景值。",
        "integrate 回归覆盖 synthetic 强制分类、缺失/非有限场景拒绝、真实正向样本晋级、过期行情阻断、CLI verify 和 API 制品读取。",
        "新闻服务：1 synthetic event / 2 versions；NVDA/TSM 证据可追溯。",
        "多因子提交：9792ed27059b1179b39cca8fca2982fe22baf86e。",
        "V4：input_mode=prod、validation_state=PASS、cvxpy=1。",
        "运行 INT-3294801BE2C27699：三路 run_backtest.py 均 PASS；V4 loader PASS；13/13 集成硬门禁通过；event study 8 行。",
        "Paper OMS：T+1_ADJ_CLOSE_PAPER，会计闭合到 0.01。",
    ])
    para(doc, "三日 synthetic 夹具产生的年化、Sharpe 和换手没有统计意义，只证明接口和执行流程。5 日主窗口仍无可用事件，当前研究决策为 INCONCLUSIVE。")
    base.add_heading(doc, "8.1 v0.5 真实数据门禁", level=2)
    bullets(doc, [
        "readiness 重新读取 CSV/Parquet，核对内容覆盖、字段语义、非有限值、重复记录和 SHA-256。",
        "PB 必须提交真实来源摄取清单；来源、映射、规范化文件哈希全部校验，规范化哈希还要与 validation、dry run、launch bundle 四方一致。",
        "当前真实 V6.5 权重截至 2026-07-17；显式 adj_close 面板截至 2024-12-31；较新的 TDX 数据为 raw close。",
        "当前 200 个真实事件均晚于可用市场输入，且缺少稳定 novelty、PIT ticker mapping、真实 PB 和连续 20 个 Paper session。",
        "封存 readiness 为 BLOCKED：23 个硬失败、0 个主窗口样本、0 个 OOS 样本、0 个验证 Paper session。",
    ])

    base.add_heading(doc, "9. 部署")
    base.add_heading(doc, "9.1 本地只读部署", level=2)
    para(doc, "启动 serve，并把 artifact directory 指向已验证产物。服务只读取工作区文件；事件研究和独立验证分别由 /api/v1/runs/latest/event-study 与 /api/v1/runs/latest/independent-validation 提供；未知路径返回 404，POST 返回 405。")
    base.add_heading(doc, "9.2 Docker", level=2)
    para(doc, "Docker 镜像安装核心包并暴露 8080。Compose 挂载 ./artifacts，健康检查 GET /health。部署前应显式选择 integrated artifact directory，而不是把 demo 当生产结果。")
    base.add_heading(doc, "9.3 外部环境", level=2)
    para(doc, "互联网或云端上线还需要主机/云账号、DNS、TLS、镜像仓库、日志、备份、回滚负责人和网络策略。当前未提供这些资源，所以本轮完成本地部署，不声称互联网生产上线。")

    base.add_heading(doc, "10. 安全与实盘门禁")
    table(doc, ["门禁", "当前状态", "解除条件"], [
        ["新闻写操作", "禁止", "无解除需求；集成保持只读"],
        ["真实 V3 数据", "缺失", "提供本地 v3_weights 和 prices"],
        ["真实事件/OOS", "机制完成；数据缺失", "真实样本通过滚动 OOS、placebo、延迟和成本压力"],
        ["PB borrow", "BLOCKED", "真实来源摄取清单 + 四方哈希一致 + validation/dry-run/launch gate"],
        ["券商连接", "不存在", "单独安全设计、凭据和发布审批"],
        ["Live launch", "BLOCKED", "全部 P0 READY + 独立验证 + 发布授权"],
    ], [2200, 1700, 5460])

    base.add_heading(doc, "11. Loop 发布流程")
    bullets(doc, [
        "Maker 完成代码、近端测试和文档；Verifier 证据单独记录。",
        "重要结论标记事实、推断、未知和决策。",
        "最后一次编辑后只执行 VERIFY -> SEAL -> CLEANUP -> FINAL。",
        "私有 GitHub 发布、合并、标签和 v0.4 替换都必须等待同一真实运行通过全部 readiness 硬门禁。",
    ])
    return doc


def save(doc: Document, filename: str) -> Path:
    OUT.mkdir(parents=True, exist_ok=True)
    path = OUT / filename
    doc.save(path)
    return path


if __name__ == "__main__":
    results = [
        save(build_selection(), "01_开源选型与MVP建议.docx"),
        save(build_prd(), "02_产品需求文档_PRD.docx"),
        save(build_development(), "03_开发与部署文档.docx"),
    ]
    for result in results:
        print(result)
